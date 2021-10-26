#!/usr/bin/python3
import sys
import re
from docx import Document

A_CODE = ord('A')
SEP = '\t'
REVERSE = '-r'

if __name__ == '__main__':
    if len(sys.argv) == 1:
        print("Pass a .docx file. Use the `%s' switch after the file name to print answer first." % REVERSE)
        exit()
        
    document = Document(sys.argv[1])
    reverse = len(sys.argv) > 2 and sys.argv[2] == REVERSE

    num_q = len(document.tables)
    answers = []

    for table in document.tables:
        ans = [None] * int(len(table.rows) * (len(table.columns) / 2))

        for n in range(1, len(table.columns), 2):
            for o in range(len(table.columns[n].cells)):
                ans[ord(table.columns[n - 1].cells[o].text[0]) - A_CODE] = table.columns[n].cells[o].text

        answers.append(ans)

    q_and_a = [[None, None] for _ in range(num_q)]

    q_re = re.compile(r'(^[0-9]{1,2})(?:\.\W*)(.+$)')
    a_re = re.compile(r'(?:^ANS:\W*)(\w)')
    last_q = -1

    for p in document.paragraphs:
        ptext = p.text.strip()
        if len(ptext) > 0:
            q_match = q_re.search(ptext)
            a_match = a_re.search(ptext)
            if q_match:
                last_q = int(q_match.group(1)) - 1
                q_and_a[last_q][0] = q_match.group(2)
            if a_match:
                q_and_a[last_q][1] = answers[last_q][ord(a_match.group(1)) - A_CODE]

    for q in q_and_a:
        print(q[reverse], q[not reverse], sep=SEP)
